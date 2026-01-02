"""
多模态融合理解服务
支持图像、文档、文本、语音的融合理解
"""
import logging
import base64
import io
from typing import Dict, List, Optional, Any, Union
from pathlib import Path
import re

logger = logging.getLogger(__name__)


class ImageProcessor:
    """图像处理器"""
    
    def __init__(self):
        self.supported_formats = ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp']
    
    async def process_image(
        self,
        image_data: Union[bytes, str],
        task: str = "auto",
        question: Optional[str] = None
    ) -> Dict:
        """
        处理图像（异步，支持真实多模态API）
        
        Args:
            image_data: 图像数据（字节流或base64字符串）
            task: 任务类型 (ocr/caption/qa/auto)
            question: 问题（用于视觉问答）
        
        Returns:
            处理结果
        """
        try:
            # 转换图像数据
            if isinstance(image_data, str):
                # Base64解码
                image_bytes = base64.b64decode(image_data)
            else:
                image_bytes = image_data
            
            # 根据任务类型处理（异步调用）
            if task == "ocr" or (task == "auto" and self._detect_text_in_image(image_bytes)):
                return await self._extract_text_from_image(image_bytes)
            elif task == "caption":
                return await self._generate_image_caption(image_bytes)
            elif task == "qa":
                return await self._answer_question_about_image(image_bytes, question)
            else:
                # 默认：生成描述
                return await self._generate_image_caption(image_bytes)
        except Exception as e:
            logger.error(f"图像处理失败: {e}", exc_info=True)
            return {
                "type": "error",
                "content": f"图像处理失败: {str(e)}",
                "success": False
            }
    
    def _detect_text_in_image(self, image_bytes: bytes) -> bool:
        """检测图像中是否包含文字（简化实现）"""
        # 简化实现：实际应该使用OCR模型检测
        # 这里假设所有图像都可能包含文字
        return True
    
    async def _extract_text_from_image(self, image_bytes: bytes) -> Dict:
        """从图像中提取文字（OCR，使用通义千问多模态API）"""
        try:
            # 使用多模态适配器
            from app.ai_engine.multimodaladapter import get_multimodal_adapter
            adapter = await get_multimodal_adapter()
            
            if adapter:
                result = await adapter.extract_text_from_image(image_bytes)
                return {
                    "type": "text",
                    "content": result.get("content", ""),
                    "method": "qwen-vl-ocr",
                    "success": True
                }
            else:
                # 降级到简化实现
                logger.warning("多模态适配器不可用，使用简化OCR实现")
                return {
                    "type": "text",
                    "content": "[OCR识别结果：需要配置API密钥]",
                    "method": "simplified",
                    "success": False,
                    "note": "请配置DASHSCOPE_API_KEY以使用真实OCR"
                }
        except Exception as e:
            logger.error(f"OCR识别失败: {e}", exc_info=True)
            return {
                "type": "text",
                "content": f"[OCR识别失败: {str(e)}]",
                "method": "error",
                "success": False
            }
    
    async def _generate_image_caption(self, image_bytes: bytes) -> Dict:
        """生成图像描述（使用通义千问多模态API）"""
        try:
            # 使用多模态适配器
            from app.ai_engine.multimodaladapter import get_multimodal_adapter
            adapter = await get_multimodal_adapter()
            
            if adapter:
                result = await adapter.generate_image_caption(image_bytes)
                return {
                    "type": "description",
                    "content": result.get("content", ""),
                    "method": "qwen-vl-caption",
                    "success": True
                }
            else:
                # 降级到简化实现
                logger.warning("多模态适配器不可用，使用简化图像描述实现")
                return {
                    "type": "description",
                    "content": "[图像描述：需要配置API密钥]",
                    "method": "simplified",
                    "success": False,
                    "note": "请配置DASHSCOPE_API_KEY以使用真实图像描述"
                }
        except Exception as e:
            logger.error(f"图像描述生成失败: {e}", exc_info=True)
            return {
                "type": "description",
                "content": f"[图像描述失败: {str(e)}]",
                "method": "error",
                "success": False
            }
    
    async def _answer_question_about_image(self, image_bytes: bytes, question: Optional[str] = None) -> Dict:
        """回答关于图像的问题（使用通义千问多模态API）"""
        if not question:
            question = "请描述这张图片的主要内容"
        
        try:
            # 使用多模态适配器
            from app.ai_engine.multimodaladapter import get_multimodal_adapter
            adapter = await get_multimodal_adapter()
            
            if adapter:
                result = await adapter.answer_question_about_image(image_bytes, question)
                return {
                    "type": "answer",
                    "content": result.get("content", ""),
                    "method": "qwen-vl-qa",
                    "question": question,
                    "success": True
                }
            else:
                # 降级到简化实现
                logger.warning("多模态适配器不可用，使用简化视觉问答实现")
                return {
                    "type": "answer",
                    "content": "[视觉问答结果：需要配置API密钥]",
                    "method": "simplified",
                    "question": question,
                    "success": False,
                    "note": "请配置DASHSCOPE_API_KEY以使用真实视觉问答"
                }
        except Exception as e:
            logger.error(f"视觉问答失败: {e}", exc_info=True)
            return {
                "type": "answer",
                "content": f"[视觉问答失败: {str(e)}]",
                "method": "error",
                "question": question,
                "success": False
            }


class DocumentProcessor:
    """文档处理器"""
    
    def __init__(self):
        self.supported_formats = {
            '.txt': self._parse_text,
            '.md': self._parse_markdown,
            '.json': self._parse_json,
            '.pdf': self._parse_pdf,
            '.docx': self._parse_docx,
            '.doc': self._parse_doc
        }
    
    def process_document(
        self,
        file_data: bytes,
        filename: str,
        extract_structure: bool = True
    ) -> Dict:
        """
        处理文档
        
        Args:
            file_data: 文件数据
            filename: 文件名
            extract_structure: 是否提取文档结构
        
        Returns:
            处理结果
        """
        try:
            file_ext = Path(filename).suffix.lower()
            
            if file_ext not in self.supported_formats:
                return {
                    "type": "error",
                    "content": f"不支持的文件格式: {file_ext}",
                    "success": False
                }
            
            parser = self.supported_formats[file_ext]
            result = parser(file_data, filename)
            
            if extract_structure:
                result["structure"] = self._extract_structure(result.get("text", ""))
            
            result["filename"] = filename
            result["file_type"] = file_ext
            result["success"] = True
            
            return result
        except Exception as e:
            logger.error(f"文档处理失败: {e}")
            return {
                "type": "error",
                "content": f"文档处理失败: {str(e)}",
                "success": False
            }
    
    def _parse_text(self, file_data: bytes, filename: str) -> Dict:
        """解析文本文件"""
        try:
            text = file_data.decode('utf-8')
            return {
                "type": "text",
                "text": text,
                "metadata": {"encoding": "utf-8"}
            }
        except UnicodeDecodeError:
            try:
                text = file_data.decode('gbk')
                return {
                    "type": "text",
                    "text": text,
                    "metadata": {"encoding": "gbk"}
                }
            except:
                return {
                    "type": "error",
                    "text": "",
                    "metadata": {}
                }
    
    def _parse_markdown(self, file_data: bytes, filename: str) -> Dict:
        """解析Markdown文件"""
        result = self._parse_text(file_data, filename)
        result["type"] = "markdown"
        return result
    
    def _parse_json(self, file_data: bytes, filename: str) -> Dict:
        """解析JSON文件"""
        import json
        try:
            data = json.loads(file_data.decode('utf-8'))
            return {
                "type": "json",
                "text": json.dumps(data, ensure_ascii=False, indent=2),
                "data": data,
                "metadata": {"format": "json"}
            }
        except Exception as e:
            return {
                "type": "error",
                "text": "",
                "metadata": {"error": str(e)}
            }
    
    def _parse_pdf(self, file_data: bytes, filename: str) -> Dict:
        """解析PDF文件（使用专业PDF解析库）"""
        try:
            # 尝试使用pdfplumber（推荐，功能强大）
            try:
                import pdfplumber
                import io
                
                pdf_file = io.BytesIO(file_data)
                text_parts = []
                pages_count = 0
                
                with pdfplumber.open(pdf_file) as pdf:
                    pages_count = len(pdf.pages)
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            text_parts.append(text)
                
                full_text = "\n\n".join(text_parts)
                logger.info(f"PDF解析成功: {filename}, 页数: {pages_count}, 文本长度: {len(full_text)}")
                
                return {
                    "type": "pdf",
                    "text": full_text,
                    "metadata": {"format": "pdf", "pages": pages_count, "method": "pdfplumber"}
                }
            except ImportError:
                # 尝试使用PyPDF2（备选）
                try:
                    import PyPDF2
                    import io
                    
                    pdf_file = io.BytesIO(file_data)
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    
                    text_parts = []
                    for page in pdf_reader.pages:
                        text = page.extract_text()
                        if text:
                            text_parts.append(text)
                    
                    full_text = "\n\n".join(text_parts)
                    logger.info(f"PDF解析成功（PyPDF2）: {filename}, 页数: {len(pdf_reader.pages)}")
                    
                    return {
                        "type": "pdf",
                        "text": full_text,
                        "metadata": {"format": "pdf", "pages": len(pdf_reader.pages), "method": "PyPDF2"}
                    }
                except ImportError:
                    logger.warning("PDF解析库未安装，使用简化实现。安装: pip install pdfplumber 或 pip install PyPDF2")
                    return {
                        "type": "pdf",
                        "text": "[PDF内容：需要安装PDF解析库（pdfplumber或PyPDF2）]",
                        "metadata": {"format": "pdf", "pages": 0, "method": "simplified"},
                        "note": "请安装: pip install pdfplumber"
                    }
        except Exception as e:
            logger.error(f"PDF解析失败: {e}", exc_info=True)
            return {
                "type": "error",
                "text": f"[PDF解析失败: {str(e)}]",
                "metadata": {}
            }
    
    def _parse_docx(self, file_data: bytes, filename: str) -> Dict:
        """解析Word文档（使用python-docx库）"""
        try:
            try:
                from docx import Document
                import io
                
                doc_file = io.BytesIO(file_data)
                doc = Document(doc_file)
                
                text_parts = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)
                
                # 提取表格内容
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join([cell.text for cell in row.cells])
                        if row_text.strip():
                            text_parts.append(row_text)
                
                full_text = "\n".join(text_parts)
                logger.info(f"Word文档解析成功: {filename}, 文本长度: {len(full_text)}")
                
                return {
                    "type": "docx",
                    "text": full_text,
                    "metadata": {"format": "docx", "method": "python-docx", "paragraphs": len(doc.paragraphs)}
                }
            except ImportError:
                logger.warning("python-docx未安装，使用简化实现。安装: pip install python-docx")
                return {
                    "type": "docx",
                    "text": "[DOCX内容：需要安装python-docx库]",
                    "metadata": {"format": "docx", "method": "simplified"},
                    "note": "请安装: pip install python-docx"
                }
        except Exception as e:
            logger.error(f"Word文档解析失败: {e}", exc_info=True)
            return {
                "type": "error",
                "text": f"[DOCX解析失败: {str(e)}]",
                "metadata": {}
            }
    
    def _parse_doc(self, file_data: bytes, filename: str) -> Dict:
        """解析DOC文件"""
        logger.warning("使用简化的DOC解析，建议集成专业库")
        return {
            "type": "doc",
            "text": "[DOC内容：需要集成专业DOC解析库]",
            "metadata": {"format": "doc"}
        }
    
    def _extract_structure(self, text: str) -> Dict:
        """提取文档结构"""
        structure = {
            "headings": [],
            "paragraphs": [],
            "lists": [],
            "tables": []
        }
        
        # 提取标题（简化实现）
        heading_pattern = r'^#+\s+(.+)$'
        for line in text.split('\n'):
            match = re.match(heading_pattern, line)
            if match:
                level = len(line) - len(line.lstrip('#'))
                structure["headings"].append({
                    "text": match.group(1),
                    "level": level
                })
        
        # 提取段落
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        structure["paragraphs"] = [{"text": p, "index": i} for i, p in enumerate(paragraphs[:10])]
        
        return structure


class MultimodalFusionService:
    """多模态融合服务"""
    
    def __init__(self):
        self.image_processor = ImageProcessor()
        self.document_processor = DocumentProcessor()
    
    def process_multimodal_input(
        self,
        inputs: List[Dict]
    ) -> Dict:
        """
        处理多模态输入
        
        Args:
            inputs: 输入列表，每个输入包含type和data
                [
                    {"type": "text", "data": "文本内容"},
                    {"type": "image", "data": image_bytes},
                    {"type": "document", "data": file_bytes, "filename": "doc.pdf"}
                ]
        
        Returns:
            融合后的理解结果
        """
        processed_modalities = []
        
        for input_item in inputs:
            input_type = input_item.get("type")
            input_data = input_item.get("data")
            
            if input_type == "text":
                processed = self._process_text(input_data)
            elif input_type == "image":
                task = input_item.get("task", "auto")
                processed = self.image_processor.process_image(input_data, task)
            elif input_type == "document":
                filename = input_item.get("filename", "unknown")
                processed = self.document_processor.process_document(input_data, filename)
            elif input_type == "audio":
                processed = self._process_audio(input_data)
            else:
                logger.warning(f"未知的输入类型: {input_type}")
                continue
            
            processed_modalities.append({
                "type": input_type,
                "result": processed
            })
        
        # 融合理解
        fused_result = self._fuse_modalities(processed_modalities)
        
        return fused_result
    
    def _process_text(self, text: str) -> Dict:
        """处理文本"""
        return {
            "type": "text",
            "content": text,
            "length": len(text),
            "word_count": len(text.split())
        }
    
    def _process_audio(self, audio_data: bytes) -> Dict:
        """处理音频（增强实现，使用ASR服务）"""
        try:
            from app.ai_engine.kylin_sdk.client import KylinAIClient
            import asyncio
            
            # 使用ASR服务进行语音识别
            ai_client = KylinAIClient()
            
            # 异步调用ASR
            try:
                loop = asyncio.get_event_loop()
                if loop.is_running():
                    asr_result = asyncio.create_task(ai_client.recognize_speech(audio_data))
                    asr_result = loop.run_until_complete(asr_result)
                else:
                    asr_result = loop.run_until_complete(ai_client.recognize_speech(audio_data))
            except RuntimeError:
                asr_result = asyncio.run(ai_client.recognize_speech(audio_data))
            
            if asr_result.get("success") or asr_result.get("text"):
                text = asr_result.get("text", "")
                confidence = asr_result.get("confidence", 0.0)
                duration = asr_result.get("duration", 0.0)
                
                logger.info(f"✅ 音频处理成功（ASR）: {len(text)} 字符, 置信度: {confidence:.2f}")
                
                return {
                    "type": "audio",
                    "content": text,
                    "text": text,  # 兼容字段
                    "confidence": confidence,
                    "duration": duration,
                    "method": "asr",
                    "success": True
                }
            else:
                logger.warning("ASR识别失败，使用简化实现")
                return {
                    "type": "audio",
                    "content": "[音频转文本：ASR识别失败]",
                    "duration": 0,
                    "method": "fallback",
                    "success": False
                }
        except Exception as e:
            logger.warning(f"使用ASR服务处理音频失败: {e}，使用简化实现")
            return {
                "type": "audio",
                "content": "[音频转文本：需要配置ASR服务]",
                "duration": 0,
                "method": "simplified",
                "success": False,
                "note": "请配置DASHSCOPE_API_KEY以使用真实ASR"
            }
    
    def _fuse_modalities(self, modalities: List[Dict]) -> Dict:
        """融合多模态理解结果"""
        if not modalities:
            return {
                "success": False,
                "error": "没有有效的输入"
            }
        
        # 提取各模态的内容
        text_content = []
        image_content = []
        document_content = []
        audio_content = []
        
        for modality in modalities:
            mod_type = modality["type"]
            result = modality["result"]
            
            if mod_type == "text":
                text_content.append(result.get("content", ""))
            elif mod_type == "image":
                image_content.append(result)
            elif mod_type == "document":
                document_content.append(result)
            elif mod_type == "audio":
                audio_content.append(result)
        
        # 构建融合结果
        fused = {
            "success": True,
            "modalities": {
                "text": text_content,
                "image": image_content,
                "document": document_content,
                "audio": audio_content
            },
            "summary": self._generate_summary(modalities),
            "combined_text": self._combine_text_content(text_content, document_content)
        }
        
        return fused
    
    def _generate_summary(self, modalities: List[Dict]) -> str:
        """生成多模态内容摘要"""
        parts = []
        
        for modality in modalities:
            mod_type = modality["type"]
            result = modality["result"]
            
            if mod_type == "text":
                content = result.get("content", "")
                if content:
                    parts.append(f"文本内容：{content[:100]}...")
            elif mod_type == "image":
                parts.append("包含图像内容")
            elif mod_type == "document":
                filename = result.get("filename", "未知文件")
                parts.append(f"包含文档：{filename}")
            elif mod_type == "audio":
                parts.append("包含音频内容")
        
        return "；".join(parts) if parts else "无内容"
    
    def _combine_text_content(self, text_content: List[str], document_content: List[Dict]) -> str:
        """合并所有文本内容"""
        combined = []
        
        # 添加文本内容
        for text in text_content:
            if text:
                combined.append(text)
        
        # 添加文档内容
        for doc in document_content:
            doc_text = doc.get("text", "")
            if doc_text:
                combined.append(doc_text)
        
        return "\n\n".join(combined)


# 全局多模态融合服务实例
multimodal_fusion_service = MultimodalFusionService()





