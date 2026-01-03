"""
增强文档处理器
集成 easydoc、mineru、PyPDF2、python-docx 等文档处理工具
支持 PDF、Word、Excel、HTML等多种格式
"""
import logging
from typing import Dict, Optional, List
import io
from pathlib import Path

logger = logging.getLogger(__name__)


class DocumentProcessorAdvanced:
    """增强文档处理器"""
    
    def __init__(self):
        """初始化文档处理器"""
        self.supported_formats = [
            '.pdf', '.docx', '.doc', '.txt', '.md', 
            '.html', '.htm', '.xlsx', '.xls', '.csv',
            '.json', '.xml'
        ]
        
        # 检测可用的处理工具
        self.available_tools = self._detect_available_tools()
        logger.info(f"文档处理器初始化完成，支持格式: {', '.join(self.supported_formats)}")
        logger.info(f"可用工具: {', '.join(self.available_tools.keys())}")
    
    def _detect_available_tools(self) -> Dict[str, bool]:
        """检测可用的文档处理工具"""
        tools = {}
        
        # 检测 PyPDF2
        try:
            import PyPDF2
            tools['pypdf2'] = True
            logger.debug("PyPDF2 可用")
        except ImportError:
            tools['pypdf2'] = False
            logger.debug("PyPDF2 不可用")
        
        # 检测 pdfplumber
        try:
            import pdfplumber
            tools['pdfplumber'] = True
            logger.debug("pdfplumber 可用")
        except ImportError:
            tools['pdfplumber'] = False
            logger.debug("pdfplumber 不可用")
        
        # 检测 python-docx
        try:
            import docx
            tools['python-docx'] = True
            logger.debug("python-docx 可用")
        except ImportError:
            tools['python-docx'] = False
            logger.debug("python-docx 不可用")
        
        # 检测 openpyxl
        try:
            import openpyxl
            tools['openpyxl'] = True
            logger.debug("openpyxl 可用")
        except ImportError:
            tools['openpyxl'] = False
            logger.debug("openpyxl 不可用")
        
        # 检测 beautifulsoup4
        try:
            from bs4 import BeautifulSoup
            tools['beautifulsoup4'] = True
            logger.debug("beautifulsoup4 可用")
        except ImportError:
            tools['beautifulsoup4'] = False
            logger.debug("beautifulsoup4 不可用")
        
        # 检测 mineru (PDF解析工具)
        try:
            import mineru
            tools['mineru'] = True
            logger.debug("mineru 可用")
        except ImportError:
            tools['mineru'] = False
            logger.debug("mineru 不可用，可运行: pip install mineru")
        
        # 检测 easydoc (文档处理工具)
        try:
            import easydoc
            tools['easydoc'] = True
            logger.debug("easydoc 可用")
        except ImportError:
            tools['easydoc'] = False
            logger.debug("easydoc 不可用，可运行: pip install easydoc")
        
        return tools
    
    def extract_text(
        self,
        file_data: bytes,
        filename: str,
        use_enhanced: bool = True,
        method: Optional[str] = None
    ) -> Dict:
        """
        从文件中提取文本
        
        Args:
            file_data: 文件数据
            filename: 文件名
            use_enhanced: 是否使用增强工具
            method: 指定使用的方法（如 'mineru', 'easydoc', 'pdfplumber' 等）
        
        Returns:
            提取结果字典
        """
        try:
            file_ext = Path(filename).suffix.lower()
            
            if file_ext == '.pdf':
                return self._extract_pdf(file_data, use_enhanced, method)
            elif file_ext in ['.docx', '.doc']:
                return self._extract_docx(file_data, use_enhanced)
            elif file_ext in ['.xlsx', '.xls']:
                return self._extract_excel(file_data)
            elif file_ext in ['.html', '.htm']:
                return self._extract_html(file_data)
            elif file_ext in ['.txt', '.md']:
                return self._extract_text_file(file_data)
            elif file_ext in ['.json']:
                return self._extract_json(file_data)
            elif file_ext in ['.xml']:
                return self._extract_xml(file_data)
            else:
                return {
                    "success": False,
                    "error": f"不支持的文件格式: {file_ext}",
                    "text": ""
                }
        except Exception as e:
            logger.error(f"提取文本失败: {e}", exc_info=True)
            return {
                "success": False,
                "error": str(e),
                "text": ""
            }
    
    def _extract_pdf(
        self,
        file_data: bytes,
        use_enhanced: bool = True,
        method: Optional[str] = None
    ) -> Dict:
        """提取PDF文本"""
        text = ""
        method_used = "unknown"
        
        # 优先级：mineru > easydoc > pdfplumber > PyPDF2
        
        # 1. 尝试使用 mineru（如果可用且已选择）
        if use_enhanced and (method == 'mineru' or method is None) and self.available_tools.get('mineru'):
            try:
                import mineru
                # mineru的使用方法（根据实际API调整）
                # 假设mineru提供了extract_text方法
                text = mineru.extract_text(file_data)
                method_used = "mineru"
                logger.debug(f"使用 mineru 提取PDF，长度: {len(text)}")
                return {
                    "success": True,
                    "text": text,
                    "method": method_used,
                    "metadata": {"tool": "mineru", "format": "pdf"}
                }
            except Exception as e:
                logger.warning(f"mineru提取失败: {e}")
        
        # 2. 尝试使用 easydoc（如果可用且已选择）
        if use_enhanced and (method == 'easydoc' or method is None) and self.available_tools.get('easydoc'):
            try:
                import easydoc
                # easydoc的使用方法（根据实际API调整）
                text = easydoc.parse_pdf(file_data)
                method_used = "easydoc"
                logger.debug(f"使用 easydoc 提取PDF，长度: {len(text)}")
                return {
                    "success": True,
                    "text": text,
                    "method": method_used,
                    "metadata": {"tool": "easydoc", "format": "pdf"}
                }
            except Exception as e:
                logger.warning(f"easydoc提取失败: {e}")
        
        # 3. 尝试使用 pdfplumber
        if use_enhanced and (method == 'pdfplumber' or method is None) and self.available_tools.get('pdfplumber'):
            try:
                import pdfplumber
                pages_text = []
                with pdfplumber.open(io.BytesIO(file_data)) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            pages_text.append(page_text)
                text = "\n\n".join(pages_text)
                method_used = "pdfplumber"
                logger.debug(f"使用 pdfplumber 提取PDF，长度: {len(text)}")
                return {
                    "success": True,
                    "text": text,
                    "method": method_used,
                    "metadata": {"tool": "pdfplumber", "format": "pdf", "pages": len(pages_text)}
                }
            except Exception as e:
                logger.warning(f"pdfplumber提取失败: {e}")
        
        # 4. 回退到 PyPDF2
        if self.available_tools.get('pypdf2'):
            try:
                import PyPDF2
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_data))
                pages_text = []
                for page in pdf_reader.pages:
                    pages_text.append(page.extract_text())
                text = "\n\n".join(pages_text)
                method_used = "pypdf2"
                logger.debug(f"使用 PyPDF2 提取PDF，长度: {len(text)}")
                return {
                    "success": True,
                    "text": text,
                    "method": method_used,
                    "metadata": {"tool": "pypdf2", "format": "pdf", "pages": len(pages_text)}
                }
            except Exception as e:
                logger.error(f"PyPDF2提取失败: {e}")
                return {
                    "success": False,
                    "error": f"PDF提取失败: {e}",
                    "text": ""
                }
        
        # 如果没有可用工具
        return {
            "success": False,
            "error": "没有可用的PDF处理工具，请安装: pip install pdfplumber PyPDF2",
            "text": ""
        }
    
    def _extract_docx(self, file_data: bytes, use_enhanced: bool = True) -> Dict:
        """提取Word文档文本"""
        if not self.available_tools.get('python-docx'):
            return {
                "success": False,
                "error": "python-docx未安装，请运行: pip install python-docx",
                "text": ""
            }
        
        try:
            import docx
            doc = docx.Document(io.BytesIO(file_data))
            
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # 提取表格内容
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        tables_text.append(row_text)
            
            text = "\n\n".join(paragraphs)
            if tables_text:
                text += "\n\n表格内容:\n" + "\n".join(tables_text)
            
            return {
                "success": True,
                "text": text,
                "method": "python-docx",
                "metadata": {
                    "tool": "python-docx",
                    "format": "docx",
                    "paragraphs": len(paragraphs),
                    "tables": len(doc.tables)
                }
            }
        except Exception as e:
            logger.error(f"提取Word文档失败: {e}")
            return {
                "success": False,
                "error": str(e),
                "text": ""
            }
    
    def _extract_excel(self, file_data: bytes) -> Dict:
        """提取Excel文本"""
        if not self.available_tools.get('openpyxl'):
            return {
                "success": False,
                "error": "openpyxl未安装，请运行: pip install openpyxl",
                "text": ""
            }
        
        try:
            import openpyxl
            workbook = openpyxl.load_workbook(io.BytesIO(file_data))
            
            sheets_text = []
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_text = [f"工作表: {sheet_name}"]
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                    if row_text.strip():
                        sheet_text.append(row_text)
                
                sheets_text.append("\n".join(sheet_text))
            
            text = "\n\n".join(sheets_text)
            
            return {
                "success": True,
                "text": text,
                "method": "openpyxl",
                "metadata": {
                    "tool": "openpyxl",
                    "format": "xlsx",
                    "sheets": len(workbook.sheetnames)
                }
            }
        except Exception as e:
            logger.error(f"提取Excel失败: {e}")
            return {
                "success": False,
                "error": str(e),
                "text": ""
            }
    
    def _extract_html(self, file_data: bytes) -> Dict:
        """提取HTML文本"""
        if not self.available_tools.get('beautifulsoup4'):
            # 回退到基础提取
            try:
                text = file_data.decode('utf-8')
                return {
                    "success": True,
                    "text": text,
                    "method": "basic",
                    "metadata": {"tool": "basic", "format": "html"}
                }
            except Exception as e:
                return {
                    "success": False,
                    "error": str(e),
                    "text": ""
                }
        
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(file_data.decode('utf-8'), 'html.parser')
            
            # 移除script和style标签
            for script in soup(["script", "style"]):
                script.decompose()
            
            # 获取文本
            text = soup.get_text(separator='\n', strip=True)
            
            return {
                "success": True,
                "text": text,
                "method": "beautifulsoup4",
                "metadata": {"tool": "beautifulsoup4", "format": "html"}
            }
        except Exception as e:
            logger.error(f"提取HTML失败: {e}")
            return {
                "success": False,
                "error": str(e),
                "text": ""
            }
    
    def _extract_text_file(self, file_data: bytes) -> Dict:
        """提取纯文本文件"""
        try:
            text = file_data.decode('utf-8')
            return {
                "success": True,
                "text": text,
                "method": "basic",
                "metadata": {"tool": "basic", "format": "text"}
            }
        except UnicodeDecodeError:
            try:
                text = file_data.decode('gbk')
                return {
                    "success": True,
                    "text": text,
                    "method": "basic",
                    "metadata": {"tool": "basic", "format": "text", "encoding": "gbk"}
                }
            except Exception as e:
                return {
                    "success": False,
                    "error": f"文本解码失败: {e}",
                    "text": ""
                }
    
    def _extract_json(self, file_data: bytes) -> Dict:
        """提取JSON文件"""
        try:
            import json
            data = json.loads(file_data.decode('utf-8'))
            text = json.dumps(data, ensure_ascii=False, indent=2)
            return {
                "success": True,
                "text": text,
                "method": "json",
                "metadata": {"tool": "json", "format": "json"}
            }
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "text": ""
            }
    
    def _extract_xml(self, file_data: bytes) -> Dict:
        """提取XML文件"""
        try:
            text = file_data.decode('utf-8')
            return {
                "success": True,
                "text": text,
                "method": "basic",
                "metadata": {"tool": "basic", "format": "xml"}
            }
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "text": ""
            }


# 全局文档处理器实例
document_processor_advanced = DocumentProcessorAdvanced()

