#!/usr/bin/env python3
"""生成知弈 AgentOS Harness 工程技术报告 Word 文档"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import datetime

doc = Document()

# ── 全局样式 ──
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.color.rgb = RGBColor(0x1a, 0x3c, 0x34)
    heading_style.font.name = 'Microsoft YaHei'
    heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

def add_code_block(doc, code_text):
    """添加代码块"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2d, 0x2d, 0x2d)
    p.style = doc.styles['Normal']
    return p

def add_table_note(doc, text):
    """添加表注"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.italic = True

# ═══════════════════════════════════════════════════════════════
# 封面
# ═══════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('知弈 AgentOS\nHarness 工程技术报告')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x3c, 0x34)

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('面向超长程复杂任务的动态异构群体智能架构\n——运行时内核、执行引擎与治理体系设计')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x4a, 0x6b, 0x63)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run(f'版本：V1.0\n日期：{datetime.date.today().strftime("%Y年%m月%d日")}')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 目录占位
# ═══════════════════════════════════════════════════════════════
doc.add_heading('目  录', level=1)
doc.add_paragraph('（生成后请在 Word 中右键更新目录域）')
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 第一章 项目概述
# ═══════════════════════════════════════════════════════════════
doc.add_heading('第一章  项目概述', level=1)

doc.add_heading('1.1 项目定位', level=2)
doc.add_paragraph(
    '知弈 AgentOS 是一个面向职业任务的智能体运行时系统。其核心使命不是构建又一个'
    '大模型聊天应用，而是将合同审查、法律分析、教学设计、需求分析等专业任务，'
    '纳入 Task、Workflow、Trace、Review、Checkpoint 和 Artifact 的可治理生命周期。'
    '系统通过 Harness 工程——即运行时内核与治理体系——把大模型的"黑盒文本续写"'
    '转变为可调度、可追踪、可审核、可恢复的结构化工程流程。'
)

doc.add_heading('1.2 核心问题', level=2)
doc.add_paragraph(
    '当前大模型应用普遍存在以下结构性问题：'
)
problems = [
    '过程不可见：用户输入直接交给模型续写，中间推理过程不可拆解、不可审核。',
    '状态不可恢复：任务一旦失败必须从头开始，消耗大量算力和时间成本。',
    '多智能体通信高熵：Agent 之间通过自然语言自由对话，Token 冗余爆炸、噪声级联放大。',
    '协作拓扑静态化：多 Agent 协作采用固定角色编排，无法根据任务语义动态调整。',
    '缺乏工程化治理：没有统一的 Trace、Review、Checkpoint 机制，无法满足专业场景的合规要求。',
]
for p_text in problems:
    doc.add_paragraph(p_text, style='List Bullet')

doc.add_heading('1.3 总体技术路线', level=2)
doc.add_paragraph(
    '系统的技术路线遵循"内核-应用分离、治理-执行解耦、静态-动态互补"三大原则：'
)
doc.add_paragraph(
    '第一，运行时内核（agentOS/）与行业应用（agent/packs/）严格分层。内核只定义通用接口、'
    '注册机制、调度器、治理设施和执行适配器协议，不承载任何具体行业逻辑。法律、教育、'
    '编程、写作等领域的 Agent、Skill、Workflow 和 Prompt 全部位于独立的领域包（Pack）中，'
    '通过注册机制注入内核。这保证了内核的通用性和可扩展性。'
)
doc.add_paragraph(
    '第二，治理体系（Trace/Review/Checkpoint/Evaluation）与执行引擎（native/acg/langgraph）'
    '完全解耦。三种执行引擎共享同一套治理接口，每种引擎只需实现 ExecutionAdapter 协议'
    '的 start() 和 apply_review() 两个方法。这使得执行引擎可以独立演进，治理能力可以'
    '统一升级，互不干扰。'
)
doc.add_paragraph(
    '第三，采用"静态优先，动态补位"的混合规划策略。对于已知任务类型，系统复用经过验证的'
    '工作流模板（静态优选，零规划开销）；对于未知或复杂任务，系统通过认知规划器动态构建'
    '智能体计算图（ACG），生成完整的节点-边-约束蓝图（动态补位）。这种策略同时保证了'
    '已知任务的质量保障和未知任务的适应能力。'
)

doc.add_heading('1.4 系统架构总览', level=2)
doc.add_paragraph(
    '系统采用四层架构：前端工作台（Vue 3 + TypeScript）、Java 网关层（Spring Boot + JWT）、'
    'Python 应用服务层（FastAPI + LangGraph 适配器）、AgentOS 内核层（WorkflowRuntime + '
    'ACG 引擎 + 规划器 + 治理设施）。其中 AgentOS 内核层是 Harness 工程的核心所在。'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 第二章 Harness 工程概述
# ═══════════════════════════════════════════════════════════════
doc.add_heading('第二章  Harness 工程概述', level=1)

doc.add_heading('2.1 Harness 的定义与内涵', level=2)
doc.add_paragraph(
    '在本项目中，Harness（运行时挽具）指的是 AgentOS Core Runtime —— 一套完整的'
    '智能体任务生命周期治理框架。它不等同于某个执行引擎（如 LangGraph），也不等同于'
    '某个模型调用网关（LLM Gateway），而是一个更高层次的抽象：它是所有 Task、Workflow、'
    '执行引擎、治理策略和状态恢复的"中央调度与协调层"。'
)
doc.add_paragraph(
    '如果将大模型应用比作一辆汽车，大模型本身就是发动机（提供动力），而 Harness 则是'
    '整车的底盘、变速箱、制动系统和仪表盘——它决定了动力如何传递、何时换挡、如何制动、'
    '驾驶状态如何呈现。没有 Harness，大模型只是一台裸机；有了 Harness，它才成为一台'
    '可驾驭的机器。'
)

doc.add_heading('2.2 Harness 的核心模块', level=2)
doc.add_paragraph(
    'Harness 由六大核心模块组成，每个模块承担明确的职责：'
)

modules = [
    ('WorkflowRuntime（总调度器）',
     '创建和管理 Task 对象、绑定 Workflow、选择执行适配器、驱动运行状态转换、'
     '处理审核决策、管理恢复流程。它是 Harness 的"大脑"，所有其它模块都通过它'
     '进行协调。代码入口：agentOS/src/agentos/core/runtime.py'),
    ('Execution Adapters（执行适配器）',
     '定义 ExecutionAdapter 协议，支持三种执行引擎并存：Native（线性逐步执行）、'
     'ACG（就绪集并行调度）、LangGraph（外部 StateGraph 桥接）。引擎选择由 '
     'WorkflowDefinition.runtimeEngine 字段决定。代码入口：agentOS/src/agentos/core/execution/'),
    ('Governance Layer（治理层）',
     '包含 TraceStore（18 种事件类型的全链路记录）、CheckpointStore（步骤级状态快照'
     '与恢复）、ReviewManager（人工审核决策管理）、WorkflowEvaluator（聚合指标计算）。'
     '这是 Harness 的"可观测性 + 可恢复性"基础设施。代码入口：agentOS/src/agentos/core/governance/'),
    ('ACG Engine（智能体计算图引擎）',
     '包含 ACGBlueprint（6 种节点 + 7 种边类型的 DAG 图模型）、图算法（环检测、拓扑排序、'
     '就绪集计算）、线性升格（promote）机制、以及就绪集并行调度执行器（ACGExecutor）。'
     '这是 Harness 最具创新性的部分。代码入口：agentOS/src/agentos/core/acg/'),
    ('Communication Layer（低熵通信层）',
     'ContextAssembler 按 input_spec 契约精准装配下游上下文（按字段提取而非全量倾倒）、'
     'ProvenanceLedger 记录数据血缘（谁消费了谁的哪些字段）、Token 估算与节省率度量。'
     '代码入口：agentOS/src/agentos/core/communication/'),
    ('Cognitive Planning Engine（认知规划引擎）',
     '"静态优先，动态补位"混合策略。IntentParser 解析意图、TemplateMatcher 匹配既有模板、'
     'CognitiveRouter 完成能力-智能体绑定、ACGBuilder 动态构建 ACG 蓝图。'
     '代码入口：agentOS/src/agentos/core/planning/'),
]
for title, desc in modules:
    doc.add_heading(title, level=3)
    doc.add_paragraph(desc)

doc.add_heading('2.3 Harness 的关键设计约束', level=2)
doc.add_paragraph(
    'Harness 的设计遵循以下铁律，这些约束保证了系统的架构纯净性和可演进性：'
)
constraints = [
    '内核不依赖应用：agentOS/src/agentos/core 不允许直接 import app.*、app.graphs.* 或 langgraph。',
    '引擎可互换：三种执行引擎共享同一套治理接口（Trace/Review/Checkpoint），互不干扰。',
    '行业不进内核：法律、教育等具体领域 Agent 和 Workflow 全部放在 agent/packs/ 中，通过注册机制注入。',
    '状态机强约束：Task 和 Step 的状态转换必须通过 StateMachine 验证，禁止非法跳转。',
    '故障注入不污染生产：故障注入通过 task.input.faultInjection 声明式配置，与正常执行逻辑完全解耦。',
]
for c in constraints:
    doc.add_paragraph(c, style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 第三章 核心设计思路与技术路线
# ═══════════════════════════════════════════════════════════════
doc.add_heading('第三章  核心设计思路与技术路线', level=1)

doc.add_heading('3.1 设计思路一：从"模型中心"到"治理中心"', level=2)
doc.add_paragraph(
    '传统大模型应用以模型调用为中心：输入 → 模型推理 → 输出。整个流程是一个黑盒，'
    '中间状态不可见、不可控、不可恢复。知弈 AgentOS 反转了这一范式：将专业任务的生命周期'
    '治理提升为系统的一等公民。模型只是执行工具，任务治理才是系统主干。'
)
doc.add_paragraph(
    '这一设计思路的具体体现：无论底层使用什么模型（DeepSeek、通义千问、Mock），'
    '无论使用什么执行引擎（native、acg、langgraph），Trace、Review 和 Checkpoint 始终'
    '以统一的结构运行。一个任务的执行轨迹不会因为换了模型而变得不可审计，也不会因为'
    '换了执行引擎而失去恢复能力。这种"治理不随执行而变"的设计，是本系统区别于所有'
    '以模型为中心的系统（如 LangChain、AutoGPT）的根本差异。'
)

doc.add_heading('3.2 设计思路二：内核-应用分离与可插拔引擎架构', level=2)
doc.add_paragraph(
    '系统采用严格的"内核-应用分离"架构。内核（agentOS/）只提供通用接口和治理机制，'
    '应用层（agent/）通过注册机制注入具体实现。执行引擎通过 ExecutionAdapter 协议实现'
    '可插拔：'
)
add_code_block(doc, '''class ExecutionAdapter(Protocol):
    """所有执行引擎必须实现此协议"""
    async def start(self, *, task: AgentTask, run: WorkflowRun,
                     workflow: WorkflowDefinition) -> WorkflowRun: ...
    async def apply_review(self, decision: ReviewDecision) -> WorkflowRun: ...

class ExecutionAdapterFactory(Protocol):
    def __call__(self, *, runtime, workflow,
                 implementation_id: str) -> ExecutionAdapter: ...''')
doc.add_paragraph(
    '当一个工作流被启动时，Runtime 根据其 runtimeEngine 字段自动选择适配器。每个适配器'
    '按 {engine}:{implementation_id} 缓存，避免重复构造。这种设计使新引擎的引入不影响'
    '既有机能，native、acg 和 langgraph 三种引擎在同一个系统中和平共存。'
)

doc.add_heading('3.3 设计思路三：静态优先、动态补位', level=2)
doc.add_paragraph(
    '规划器采用"静态优先，动态补位"的混合策略。这一设计的核心思想是：'
    '对于系统已经掌握的任务类型，直接复用经过验证的工作流模板，产生高质量的确定性结果，'
    '同时节省规划开销（零 LLM 调用）；对于未知或复杂任务，才动态构建智能体计算图。'
)
doc.add_paragraph(
    '模板匹配基于角色类型（domain）和任务类型（intent）的一级索引筛选，'
    '再用 Dice 字符二元组相似度评分。intent 精确命中保底 0.9 分，匹配阈值设为 0.85。'
    '超过阈值则"静态优选"：通过线性升格（promote）复用模板，零规划开销。'
    '低于阈值则"动态补位"：通过认知路由完成能力-智能体绑定，再由 ACG 构建器动态生成蓝图。'
)
doc.add_paragraph(
    '这一策略在保证已知任务高效率、高质量的同时，保持了对未知任务的开放适应能力。'
    '规划器同时具备 LLM 解析和启发式回退两条路径，保证在没有 API key 的环境下也能完成规划。'
)

doc.add_heading('3.4 设计思路四：低熵通信——将 Agent 对话升级为结构化数据交换', level=2)
doc.add_paragraph(
    '传统多智能体系统允许 Agent 之间进行无边界自然语言广播，这种"高熵"通信模式导致三个问题：'
    '（1）Token 冗余爆炸——每个 Agent 都把全部历史对话传给下一个，上下文随步骤线性膨胀；'
    '（2）噪声级联放大——无关信息在多个 Agent 之间反复传递，最终淹没关键信号；'
    '（3）数据血缘不清——无法回答"这个结论从何而来"。'
)
doc.add_paragraph(
    '本系统的低熵通信协议将智能体间的"自由对话"转变为沿明确数据依赖边流动的"精准数据流"。'
    '核心机制包括：'
)
doc.add_paragraph(
    '（1）工作流引擎作为唯一通信中介——单个 Step 不直接感知或呼叫其他 Step，仅与引擎交互。'
    '杜绝了智能体间未经审计的私自通信。'
)
doc.add_paragraph(
    '（2）input_spec 数据采购清单——下游 Step 通过 fields 和 from 声明自己需要哪些字段。'
    '引擎按清单精准提取，不会把上游的全部输出"倾倒"给下游。'
)
add_code_block(doc, '''# ACG 工作流定义中声明 input_spec
# clause_classify 只声明消费 contract_type, scope, payment_terms
{"stepId": "clause_classify",
 "input": {"fields": ["contract_type", "scope", "payment_terms"]},
 "nextStepId": "risk_detect"}

# risk_detect 只消费 clauses
{"stepId": "risk_detect",
 "input": {"fields": ["clauses"]},
 "nextStepId": "legal_evidence_match"}''')
doc.add_paragraph(
    '（3）量化节省率——ContextAssembler 计算 tokensAvailable（全量倾倒所需 Token）'
    '和 tokensDelivered（按需投递实际 Token），生成 savingRatio 指标。'
    '实测数据表明，单个步骤的节省率可达 73.1%，整体流程平均节省率 22.17%。'
)
doc.add_paragraph(
    '（4）数据血缘——ProvenanceLedger 记录每一步消费了上游的哪些字段，'
    '生成完整的数据流转图，支持前向追溯（"这个结论从何而来"）和后向影响分析'
    '（"这个数据用在了何处"）。'
)

doc.add_heading('3.5 设计思路五：故障注入与自愈闭环', level=2)
doc.add_paragraph(
    '传统大模型应用面对故障（模型超时、Agent 崩溃、证据检索为空）只能从零开始重新执行。'
    '这不仅浪费算力，更意味着已经完成的部分无法复用。本系统将混沌工程引入 AI Agent 领域，'
    '设计了一套完整的故障注入与自愈机制：'
)
add_code_block(doc, '''# 用户通过 task.input 声明式配置故障注入
task.input["faultInjection"] = {
    "step_id": "risk_detect",
    "fault_type": "timeout",     # timeout | crash | empty_evidence
    "max_triggers": 1,           # 触发 1 次后自愈
}

# ACGExecutor._self_heal() 自愈路径
def _self_heal(run, task, blueprint, node_id, fault):
    if step.retry_count >= 3: return False        # 循环保护
    trace(STEP_FAILED, recoverable=True)           # 故障入轨
    checkpoint = checkpoint_store.create(run)      # 保存现场
    step.status = PENDING; step.retry_count += 1   # 局部重规划
    trace(RUN_RECOVERED, strategy="local_replan")  # 恢复入轨
    return True                                    # 下一轮重新调度''')
doc.add_paragraph(
    '故障注入与正常执行逻辑完全解耦——注入通过 task.input 配置声明，'
    '只在 ACGExecutor 的 _execute_step 中通过 FaultInjector.fire() 触发。'
    '如果 max_triggers 设置为 1，故障只会被注入一次，之后自动"痊愈"。'
    '整个自愈过程生成完整的 Trace 事件（STEP_FAILED → RUN_RECOVERED），'
    '支持事后回放分析。'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 第四章 详细设计方案
# ═══════════════════════════════════════════════════════════════
doc.add_heading('第四章  详细设计方案', level=1)

doc.add_heading('4.1 WorkflowRuntime — 中央调度器', level=2)
doc.add_paragraph(
    'WorkflowRuntime 是 Harness 的中央调度器，负责协调所有子系统。'
    '它在初始化时组装 task_manager、workflow_registry、workflow_store、'
    'state_machine、orchestrator、trace_store、checkpoint_store、review_manager、'
    'evaluator 以及 execution_adapter_factories。'
)
doc.add_paragraph(
    '其核心方法是 start()，负责解析工作流、选择执行适配器并驱动执行。'
    '_workflow_adapter() 方法根据 runtimeEngine 字段路由到正确的适配器：'
)
add_code_block(doc, '''def _workflow_adapter(self, workflow: WorkflowDefinition):
    runtime_engine = workflow.effective_runtime_engine
    implementation_id = workflow.effective_implementation_id
    adapter_key = f"{runtime_engine}:{implementation_id}"

    adapter = self._runtime_adapters.get(adapter_key)
    if adapter is None:
        if runtime_engine == "native":
            adapter = NativeWorkflowAdapter(self)       # → _start_native()
        elif runtime_engine == "acg":
            adapter = ACGWorkflowAdapter(self)           # → _start_acg() + ACGExecutor
        else:
            factory = self.execution_adapter_factories.get(runtime_engine)
            if factory is None:
                raise ValueError(f"Unsupported: {runtime_engine}")
            adapter = factory(runtime=self, workflow=workflow, ...)
        self._runtime_adapters[adapter_key] = adapter    # 缓存
    return adapter''')

doc.add_heading('4.2 ACG 计算图模型 — 系统的"脊椎"', level=2)
doc.add_paragraph(
    '智能体计算图（Agentic Computation Graph, ACG）是系统最核心的数据结构。'
    '它不仅描述"哪些步骤需要执行"，还描述"谁（哪个智能体）来执行"、"需要什么记忆"、'
    '"依赖什么证据"、"数据如何流动"。ACG 是规划器的产物、执行器的输入，'
    '是整个系统的"统一计算模型"。'
)

doc.add_heading('4.2.1 节点体系（6 种类型）', level=3)
doc.add_paragraph(
    'ACG 定义 6 种节点类型，每种节点有明确的语义和用途：'
)
nodes = [
    ('StepNode', '最小执行单元。包含 step_type、goal、input_spec、output_spec、'
     'assigned_agent_id、review_required、retry_limit 等字段。执行器只调度 StepNode。'),
    ('AgentNode', '智能体能力标签。包含 role、model_name、capability_tags、max_concurrency。'
     '一个 Agent 可执行多个 Step（通过 EXECUTION 边绑定）。'),
    ('SkillNode', '工具能力描述。包含 skill_type、tool_name。'),
    ('MemoryNode', '上下文存储节点。包含 memory_type、storage_type、retention_policy。'
     'Step 通过 WRITE/READ 边与其交互。'),
    ('EvidenceNode', '审计痕迹节点。包含 evidence_type、source。'
     'Step 通过 SUPPORT 边引用证据。'),
    ('ControlNode', '流程控制节点。支持 START、END、IF、LOOP、PARALLEL、CONSENSUS。'
     '已满足前置依赖的控制节点会被执行器自动并入完成集。'),
]
for name, desc in nodes:
    p = doc.add_paragraph()
    run = p.add_run(f'{name}：')
    run.bold = True
    p.add_run(desc)

doc.add_heading('4.2.2 边体系（7 种类型）', level=3)
doc.add_paragraph(
    '边定义了节点之间的关系，但只有 DEPENDENCY 边参与执行 DAG 构建，'
    '其余边由通信器、记忆器、审计器分别消费。这种"执行先后与数据/记忆/证据关系解耦"'
    '的设计是 ACG 的核心洞察。'
)
edges = [
    ('DEPENDENCY（任务依赖）', '定义执行先后顺序，执行器据此计算就绪集。'),
    ('COMMUNICATION（数据流）', '定义数据传递关系，通信器据此装配 ContextPack。'),
    ('CONTROL_FLOW（控制流）', '定义条件分支和流程控制路径。'),
    ('EXECUTION（Agent→Step）', '定义智能体到执行步骤的绑定关系。'),
    ('WRITE / READ（Step↔Memory）', '定义记忆读写，记忆器据此管理上下文持久化。'),
    ('SUPPORT（Evidence→Step）', '定义证据支撑关系，审计器据此追溯依据链。'),
]
for name, desc in edges:
    p = doc.add_paragraph()
    run = p.add_run(f'{name}：')
    run.bold = True
    p.add_run(desc)

doc.add_heading('4.2.3 图算法', level=3)
doc.add_paragraph(
    'ACG 提供完整的图算法基础设施，所有算法仅操作 DEPENDENCY 子图：'
)
add_code_block(doc, '''def ready_steps(blueprint, completed) -> List[str]:
    """计算就绪集：所有 DEPENDENCY 前驱已完成、自身未完成的 STEP 节点。
    这是执行器并行调度的核心——返回的 step 之间彼此无依赖，可并发执行。"""
    ready = []
    for step in blueprint.step_nodes():
        if step.node_id in completed:
            continue
        deps = blueprint.dependency_sources(step.node_id)
        if all(dep in completed for dep in deps):
            ready.append(step.node_id)
    return ready

def detect_cycle(blueprint) -> List[str]:
    """三色 DFS 环检测——规划器交付前验证图的非循环性。"""

def topological_order(blueprint) -> List[str]:
    """Kahn 算法拓扑排序——验证 DAG 合法性。"""

def validate_blueprint(blueprint):
    """悬空依赖检查 + 环检测——规划器交付前的图级验证。"""''')

doc.add_heading('4.3 线性升格 — 存量工作流零改动接入', level=2)
doc.add_paragraph(
    'promote_workflow_to_acg() 是本系统实现"向下兼容"的核心机制。'
    '它把一条基于 nextStepId 的线性步骤链，无损转换为一张 DEPENDENCY 边串联的 DAG。'
    '转换后的图可被 ACG 执行器以"就绪集调度"方式执行，行为与原线性执行完全一致——'
    '存量工作流零代码改动接入新架构。'
)
add_code_block(doc, '''def promote_workflow_to_acg(workflow, *, task_id=None, enrich=True):
    """把线性 WorkflowDefinition 升格为 ACGBlueprint。"""
    steps = list(workflow.steps)
    blueprint = ACGBlueprint(...)

    # 1) 每个 WorkflowStep → StepNode
    for definition in steps:
        node = StepNode(nodeId=definition.step_id, name=definition.name, ...)
        blueprint.nodes.append(node)

    # 2) 相邻 Step 之间连 DEPENDENCY 边
    for index, definition in enumerate(steps):
        target_id = definition.next_step_id or steps[index+1].step_id
        blueprint.edges.append(ACGEdge(
            sourceId=definition.step_id, targetId=target_id,
            edgeType=EdgeType.DEPENDENCY))

    # 3) enrich=True 时注入认知节点
    if enrich:
        _inject_cognitive_nodes(blueprint, steps)
    return blueprint''')
doc.add_paragraph(
    'enrich=True（默认）时，升格过程还会自动注入认知协作节点：为每个 Step 创建同名 '
    'AgentNode 并通过 EXECUTION 边绑定，为产出结论的 Step 注入 MemoryNode 和 WRITE 边，'
    '为需要外部依据的 Step 注入 EvidenceNode 和 SUPPORT 边。这些注入节点不参与执行调度，'
    '仅丰富拓扑的认知协作语义和可视化表达。'
)

doc.add_heading('4.4 ACG 执行器 — 就绪集并行调度', level=2)
doc.add_paragraph(
    'ACGExecutor 是就绪集并行调度引擎，借鉴了操作系统内核的进程调度思想：'
    '以 Step 为基本调度单元，将执行模型从"链表单指针顺序推进"升级为'
    '"DAG 就绪集并行驱动"。核心循环 _drive() 的伪代码如下：'
)
add_code_block(doc, '''async def _drive(self, *, task, run, workflow, blueprint):
    completed = set(run.completed_step_ids)
    while True:
        # 1) 处理控制节点（START/PARALLEL/CONSENSUS）
        _resolve_control_nodes(blueprint, completed)

        # 2) 计算就绪集并过滤
        ready = _eligible_steps(blueprint, run, completed)

        # 3) 无就绪步骤 -> 完成或等待审核
        if not ready:
            if _all_steps_done(): complete_run(); return
            return  # 等待人审

        # 4) 并行执行本批就绪步骤
        batch = ready[:max_parallelism]
        results = await asyncio.gather(*[
            _execute_step(task, run, workflow, blueprint, nid)
            for nid in batch
        ], return_exceptions=True)

        # 5) 结算：完成/失败/等待审核/自愈
        waiting_review = False
        for node_id, outcome in zip(batch, results):
            if isinstance(outcome, InjectedFault):
                if _self_heal(run, task, blueprint, node_id, outcome):
                    continue       # 自愈续跑
                mark_failed(); return
            if outcome == WAITING_REVIEW: waiting_review = True
            elif outcome == COMPLETED: completed.add(node_id)

        if waiting_review: return  # 人审中断，等待 apply_review''')

doc.add_paragraph(
    '关键特性：线性工作流升格后每轮就绪集恰好一个节点，等效于原线性行为；'
    '如果蓝图中有并行分支（如 A→{B,C}→D），则 B 和 C 将在同一轮并发执行。'
    '这是"动态拓扑可见性"的直接体现——执行顺序由图的拓扑结构决定，'
    '而非预设的链表。'
)

doc.add_heading('4.5 低熵通信 — ContextAssembler', level=2)
doc.add_paragraph(
    'ContextAssembler 是低熵通信协议的运行时核心。它在每个 Step 执行前，'
    '根据该 Step 的 input_spec 契约，从上游 Step 的输出中精准提取所需字段，'
    '组装成 ContextPack 传递给下游 Agent。核心逻辑如下：'
)
add_code_block(doc, '''def assemble(self, *, blueprint, step_node, upstream_outputs):
    # 1) 确定上游来源
    dep_sources = blueprint.dependency_sources(step_node.node_id)
    source_ids = [sid for sid in dep_sources if sid in upstream_outputs]

    # 2) 计算可获取 token 总量（全量倾倒时的消耗）
    tokens_available = sum(estimate_tokens(upstream_outputs[s]) for s in source_ids)

    # 3) 按 input_spec 精准提取
    spec = step_node.input_spec or {}
    if spec.get("from"):     # 定向提取：{stepId: [field, ...]}
        delivered = pick_by_source_map(spec["from"])
    elif spec.get("fields"): # 字段清单提取：[field1, field2, ...]
        delivered = pick_fields(spec["fields"])
    else:                    # 回退：无清单则透传全部（兼容旧行为）
        delivered = passthrough(source_ids)

    # 4) 聚合证据链
    evidence_refs = aggregate_evidence(source_ids)

    # 5) 低熵度量
    tokens_delivered = estimate_tokens(delivered)
    saving_ratio = 1 - tokens_delivered / tokens_available

    # 6) 血缘记账
    ledger.record_consumption(step_node, source_ids, consumed_fields)
    return ContextPack(data=delivered, evidence_refs=evidence_refs,
                        saving_ratio=saving_ratio, ...)''')

doc.add_paragraph(
    '实测数据：在合同审查 6 步流程中，低熵通信累计节省 203 个 Token（22.17%），'
    '单个步骤最高节省率 73.1%（clause_classify 只从 contract_parse 的 9 个字段中'
    '提取 3 个所需字段，其余 6 个字段完全不传递）。'
)

doc.add_heading('4.6 规划器 — 从意图到蓝图', level=2)
doc.add_paragraph(
    '认知规划引擎（PlanningEngine）采用"静态优先，动态补位"混合策略，'
    '将用户自然语言意图转化为可执行的 ACGBlueprint。其 plan() 方法是整个规划流程的入口：'
)
add_code_block(doc, '''def plan(self, *, task_id, intent, domain, task_type):
    # 1) 意图解析 → TaskSemanticProfile
    profile = self.intent_parser.parse(intent=intent, domain=domain)

    # 2) 静态优先 — 模板匹配
    match = self.template_matcher.match(profile)
    if self.template_matcher.is_hit(match):
        blueprint = promote_workflow_to_acg(match.workflow)
        return PlanResult(blueprint, strategy="static_template", ...)

    # 3) 动态补位 — 认知路由 + ACG 构建
    network = self.cognitive_router.route(profile, domain=domain)
    blueprint = self.acg_builder.build(task_id, profile, network)
    return PlanResult(blueprint, strategy="dynamic_generation", ...)''')
doc.add_paragraph(
    '意图解析采用分层设计：Core 定义 IntentLLM 协议，app 层注入真实 DeepSeek 网关；'
    '未注入或调用失败时回落确定性启发式（关键词→能力映射 + 长度→复杂度），'
    '保证 Core 离线可测、规划不被 LLM 故障阻断。模板匹配使用 Dice 字符二元组相似度，'
    '无需向量数据库即可完成匹配。'
)

doc.add_heading('4.7 治理体系 — Trace / Checkpoint / Review', level=2)
doc.add_paragraph(
    '治理体系是 Harness 可观测性和可控性的基础，由三个子系统组成：'
)

doc.add_heading('4.7.1 TraceStore — 全链路事件记录', level=3)
doc.add_paragraph(
    'TraceStore 记录 18 种事件类型，覆盖任务完整生命周期：TASK_CREATED → '
    'RUN_STARTED → STEP_SCHEDULED → STEP_STARTED → DATA_CONSUMED → '
    'AGENT_CALLED → DATA_PRODUCED → STEP_SUCCEEDED → CHECKPOINT_CREATED → '
    'RUN_COMPLETED。每个事件包含 step_id、agent_name、duration_ms 和 payload。'
    '支持导出为 JSON 或 Markdown 格式。'
)

doc.add_heading('4.7.2 CheckpointStore — 步骤级快照与恢复', level=3)
doc.add_paragraph(
    '每次 Step 完成（或进入等待审核）时，系统自动创建检查点，保存整个 WorkflowRun '
    '的状态快照（steps 状态、输出、input）。恢复时，从检查点加载快照，找到第一个 '
    'PENDING/FAILED 步骤，重置为 RETRYING，重新进入执行循环。'
)
add_code_block(doc, '''async def resume_from_checkpoint(self, *, run_id, checkpoint_id):
    run = self.workflow_store.get_run(run_id)
    checkpoint = self.checkpoint_store.find(run, checkpoint_id)
    # 恢复步骤状态
    snapshot_steps = checkpoint.state_snapshot.get("steps", [])
    if snapshot_steps:
        run.steps = [WorkflowStep.model_validate(s) for s in snapshot_steps]
    # 找到第一个待恢复的步骤
    run.current_step_id = self._next_pending_step_id(run)
    run.recovery_count += 1
    # 记录恢复事件
    self.trace_store.append(run, event_type=TraceEventType.RUN_RECOVERED, ...)
    return await self._run_until_blocked(task, run, workflow)''')

doc.add_heading('4.7.3 ReviewManager — 人工审核门控', level=3)
doc.add_paragraph(
    '当 Step 声明 review_required=true 且 review_mode != "auto" 时，'
    '流程在步骤完成后暂停于 WAITING_REVIEW 状态。审核决策有五种：'
    'APPROVED（通过，继续下一步）、REJECTED（拒绝，运行失败）、RERUN（重试当前步骤）、'
    'NEED_MORE_INFO（保持等待）、CANCELLED（取消整个运行）。'
    '审核决策被记录为 REVIEW_DECIDED Trace 事件，与对应 Step 和 Run 关联。'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 第五章 创新与特色
# ═══════════════════════════════════════════════════════════════
doc.add_heading('第五章  创新与特色', level=1)

doc.add_heading('5.1 架构创新：ACG + 低熵通信 — 多智能体的结构化对话协议', level=2)
doc.add_paragraph(
    '传统多智能体框架中，Agent 之间通过自然语言全量对话通信，导致 Token 浪费、'
    '上下文污染和数据血缘不清。本系统将多智能体协作建模为有类型的计算图（ACG），'
    '定义了 6 种节点类型和 7 种边类型来承载不同语义的协作关系。COMMUNICATION 边'
    '携带 input_spec 契约声明——下游通过 fields 和 from 明确列出自己需要哪些字段，'
    'ContextAssembler 在运行时按契约精准提取投递，而非将上游全部输出倾倒给下游。'
)
doc.add_paragraph(
    '这相当于在多智能体系统中引入了"类型系统"——在自由流动的信息中引入了结构和约束，'
    '而正是这些约束让系统变得可预测、可优化、可治理。类比 HTTP 的 RESTful API 之于'
    '原始 TCP 字节流：ACG 的 COMMUNICATION 契约就是多智能体之间的 API 接口定义。'
)
doc.add_paragraph(
    '同时，ProvenanceLedger 记录每一步消费了上游的哪些字段，生成完整的数据血缘图谱。'
    '系统支持前向追溯（"这个结论从何而来"）和后向影响分析（"这个数据用在了何处"），'
    '使多智能体协作过程从黑盒走向白盒。在合同审查实测中，低熵通信实现平均节省率 22.17%，'
    '单步最高节省率 73.1%，有效验证了该方案在多步骤任务中的 Token 节约效果。'
)

doc.add_heading('5.2 技术创新：线性工作流零成本升格 — 从链表到 DAG 的自动编译', level=2)
doc.add_paragraph(
    '本系统设计了 promote_workflow_to_acg() 机制，将基于 nextStepId 的线性步骤链'
    '无损转换为 DEPENDENCY 边串联的 DAG。这是"静态优选"一侧的落地基础——保证存量'
    '工作流零代码改动即可接入 ACG 并行调度引擎。升格过程中自动注入 Agent/Memory/Evidence'
    '认知节点（不参与执行调度，只丰富拓扑语义和可视化表达），将线性链条变成多层认知网络。'
)
doc.add_paragraph(
    '该设计的创新价值在于：它类似于编译器的"中间表示提升（IR lifting）"——把低级线性'
    '代码提升为高级图 IR，然后可以在图 IR 上做并行化优化、低熵注入和可视化。这使系统'
    '具备了"渐进增强"的能力：存量资产不受影响，新能力通过图优化自然叠加。'
)

doc.add_heading('5.3 工程创新：故障注入与自愈闭环 — AI Agent 混沌工程', level=2)
doc.add_paragraph(
    '传统 AI 系统中，异常处理通常是事后补救式的——出错后查看日志、手动修复、重新运行。'
    '本系统将混沌工程的思想引入 AI Agent 领域，设计了声明式故障注入（通过 task.input '
    '配置，与生产逻辑完全解耦）和自动化自愈闭环（创建检查点 → 局部重规划 → 重新调度）。'
    '支持三种故障类型（timeout/crash/empty_evidence）、可配置触发次数（max_triggers），'
    '自愈过程全程记录 Trace 事件，支持事后回放分析。'
)
doc.add_paragraph(
    '该机制的价值不仅是测试用途——它本质上是为生产级 AI 系统提供了"韧性基础设施"。'
    '在真实的产业场景中，模型服务不稳定、外部 API 超时、检索结果为空都是常态。'
    '通过故障注入验证和自愈机制保障，系统能够在无人干预下自主从这些可恢复故障中恢复，'
    '显著提升长程任务的成功率和 ROI。'
)

doc.add_heading('5.4 设计创新：内核-应用分层架构与可插拔执行引擎', level=2)
doc.add_paragraph(
    '系统采用严格的四层架构分离：前端（Vue 3）→ Java 网关（Spring Boot + JWT）→ '
    'Python 应用层（FastAPI + 行业 Pack）→ AgentOS 内核（WorkflowRuntime + ACG + 治理）。'
    '内核不允许引用任何具体行业逻辑（langgraph、legal pack 等全部在应用层），'
    '执行引擎通过 ExecutionAdapter 协议实现可插拔（native、acg、langgraph 三引擎并存），'
    '领域能力通过 Pack Manifest 和 Registry 注入运行时。'
)
doc.add_paragraph(
    '这一架构的核心价值在于"通用性"——同一套内核可以支撑法律、教育、编程、写作等多个'
    '不同领域，新增行业能力时只需新增 agent/packs/<pack_id>/ 目录，无需修改内核代码。'
    '这使系统具备了从"法律合同审查系统"演进为"通用职业任务智能体运行时平台"的架构潜力。'
)

doc.add_heading('5.5 应用创新：JIT 编译式混合规划策略', level=2)
doc.add_paragraph(
    '系统的认知规划引擎采用"静态优先，动态补位"策略，这在 AI 工作流规划领域具有原创性。'
    '其设计思想类似于 JIT 编译器：热点方法走编译缓存（已知任务走模板复用），冷方法走'
    '解释执行（未知任务走动态生成）。模板匹配使用 Dice 字符二元组相似度，无需向量数据库'
    '即可完成匹配；同时引入能力命中加成机制（每命中一个 required_capability +0.1 分），'
    '确保"风险审查"不会因文本巧合被误匹配到"风险投资"。规划器还具备 LLM 解析和启发式'
    '回退双路径，保证在无 API key 环境下也能正常工作。'
)

doc.add_heading('5.6 应用创新：联邦 RAG 优化 — 从模型权重扩展到检索参数', level=2)
doc.add_paragraph(
    '传统联邦学习只优化模型权重。本系统将联邦学习的优化对象扩展到 RAG 检索超参数领域——'
    '每个客户端上传的不是梯度而是检索统计量（查询量、平均检索时间、最优 top_k、检索成功率、'
    '查询模式分布），服务器聚合这些统计量后生成全局最优 RAG 配置（top_k、similarity_threshold、'
    'reranking_strategy、query_expansion）。整个过程不传输原始文档或查询内容，只传输聚合统计量，'
    '隐私保护强度高于传统联邦学习。联邦优化结果通过 FederatedAdapter 注入 ACG 治理层——'
    '联邦节点越多，系统 confidence 越高，risk_adjustment 越低，直接影响人工审核的严格程度。'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 第六章 系统性能分析
# ═══════════════════════════════════════════════════════════════
doc.add_heading('第六章  系统性能分析', level=1)

doc.add_heading('6.1 图算法复杂度', level=2)
doc.add_paragraph(
    '记 ACG 节点数 N、依赖边数 E、最大并行度 P、关键路径长度 L。'
)
complexity_data = [
    ('环检测（DFS 三色）', 'O(N+E)', '规划交付前一次性验证'),
    ('拓扑排序（Kahn）', 'O(N+E)', '验证 DAG 合法性'),
    ('单轮就绪集计算', 'O(N·d̄)', 'd̄ 为平均入度，遍历未完成 Step'),
    ('全程调度轮数', 'O(L)', '每轮推进至少一层；线性图 L=N'),
    ('上下文装配（每步）', 'O(F)', 'F 为下游 input_spec 字段数，非上游全量'),
    ('血缘前向追溯', 'O(C)', 'C 为消费事件数'),
]
table = doc.add_table(rows=len(complexity_data)+1, cols=3, style='Light Grid Accent 1')
hdr = table.rows[0].cells
hdr[0].text = '操作'
hdr[1].text = '时间复杂度'
hdr[2].text = '说明'
for i, (op, time, note) in enumerate(complexity_data, 1):
    table.rows[i].cells[0].text = op
    table.rows[i].cells[1].text = time
    table.rows[i].cells[2].text = note

doc.add_heading('6.2 并行与低熵收益分析', level=2)
doc.add_paragraph(
    '并行收益：线性执行总时长 ≈ Σ 各步耗时；就绪集并行后 ≈ 关键路径耗时，'
    '理论加速比上界为"总工作量 / 关键路径"，受最大并行度 P 截断。'
)
doc.add_paragraph(
    '低熵收益：传统全量拼接下游上下文规模随步数线性膨胀 O(ΣSᵢ)；按需投递后'
    '每步仅 O(F) ≪ O(ΣSᵢ)，避免了长程任务中的 Token 冗余爆炸。'
    '在合同审查 6 步实测中，低熵通信将 Token 投递量从 1,202 降至 999，节约 203 Token，'
    '平均节省率 22.17%。在实际业务中，随着步骤数增加和输出规模增长，节省效果会更加显著。'
)

doc.add_heading('6.3 测试覆盖', level=2)
doc.add_paragraph(
    '当前测试基线：agentOS 48 项 + agent 113 项，共 161 项全绿。覆盖 ACG 模型/升格/'
    '图算法、通信协议、规划器、执行器、低熵通信、自愈机制、API 集成和既有核心链路。'
    '关键端到端验证包括：菱形图并行执行、人审中断 approve 续跑、低熵节省率度量、'
    '三类故障自愈至完成、/acg 端点暴露拓扑+血缘+恢复+低熵指标、真实 DeepSeek '
    '意图解析→跨领域动态生成 ACG。'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 第七章 总结
# ═══════════════════════════════════════════════════════════════
doc.add_heading('第七章  总结', level=1)

doc.add_paragraph(
    '知弈 AgentOS 通过 Harness 工程技术，构建了一套面向超长程复杂任务的动态异构群体智能'
    '运行时架构。系统的核心贡献在于：'
)

contributions = [
    '将大模型应用从"黑盒文本续写"升级为可调度、可追踪、可审核、可恢复的结构化工程流程——'
    '这是从"模型中心"到"治理中心"的范式转变。',
    '通过 ACG 计算图模型和低熵通信协议，解决了多智能体系统中 Token 冗余爆炸、'
    '噪声级联放大和数据血缘不清三大核心问题。',
    '通过线性升格机制实现存量工作流的零改动兼容，通过执行适配器协议实现多引擎并存，'
    '通过内核-应用分层架构保证系统的通用性和可扩展性。',
    '通过故障注入与自愈闭环，将混沌工程引入 AI Agent 领域，为生产级 AI 系统提供了韧性保障。',
    '通过"静态优先，动态补位"的 JIT 编译式规划策略，在保证已知任务高效率、高质量的同时，'
    '保持了对未知任务的开放适应能力。',
]
for i, c in enumerate(contributions, 1):
    doc.add_paragraph(f'{i}. {c}')

doc.add_paragraph()
doc.add_paragraph(
    '系统的技术路线完整覆盖了赛题要求的超长程上下文连续性、动态异构拓扑、低熵通信、'
    '端边云调度和自愈闭环五大能力维度，并通过 161 项测试验证了核心链路的正确性。'
    '在架构层面，系统的内核-应用分离、治理-执行解耦和可插拔引擎设计，为后续演进为'
    '通用职业任务智能体运行时平台奠定了坚实的基础。'
)

# ═══════════════════════════════════════════════════════════════
# 保存
# ═══════════════════════════════════════════════════════════════
output_path = 'output/知弈AgentOS_Harness工程技术报告.docx'
doc.save(output_path)
print(f'报告已生成：{output_path}')
